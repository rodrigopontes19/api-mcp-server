from pathlib import Path
import sys
import shutil
import subprocess

TEXT_EXTS = {".txt", ".md", ".json", ".csv", ".xml", ".yaml", ".yml", ".log", ".ini"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}

OCR_LANG = "por+eng"


def run_tesseract(path: Path) -> str:
    if not shutil.which("tesseract"):
        raise RuntimeError("Tesseract não está no PATH.")
    result = subprocess.run(
        ["tesseract", str(path), "stdout", "-l", OCR_LANG, "--psm", "6"],
        capture_output=True,
        text=True
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "OCR falhou.")
    return result.stdout


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for i, table in enumerate(doc.tables, start=1):
        parts.append(f"\n[TABELA {i}]")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
            parts.append(row_text)

    return "\n".join(parts).strip()


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    import fitz

    reader = PdfReader(str(path))
    doc = fitz.open(str(path))
    pages = []

    total_pages = min(len(reader.pages), len(doc))
    for idx in range(total_pages):
        page_no = idx + 1
        text = ""
        try:
            text = (reader.pages[idx].extract_text() or "").strip()
        except Exception:
            text = ""

        if len(text) < 40:
            if shutil.which("tesseract"):
                pix = doc[idx].get_pixmap(dpi=200)
                tmp = Path("_extracted") / f"__ocr_{path.stem}_{page_no}.png"
                pix.save(tmp)
                try:
                    text = run_tesseract(tmp).strip()
                finally:
                    if tmp.exists():
                        tmp.unlink()
            else:
                try:
                    text = (doc[idx].get_text("text") or "").strip()
                except Exception:
                    text = ""

        pages.append(f"===== PÁGINA {page_no} =====\n{text}")

    return "\n\n".join(pages).strip()


def read_image(path: Path) -> str:
    return run_tesseract(path).strip()


def main():
    if len(sys.argv) < 2:
        raise SystemExit("Uso: python tools/extract_doc.py <arquivo>")

    source = Path(sys.argv[1]).resolve()
    if not source.exists():
        raise SystemExit(f"Arquivo não encontrado: {source}")

    ext = source.suffix.lower()

    if ext in TEXT_EXTS:
        content = read_text_file(source)
    elif ext == ".docx":
        content = read_docx(source)
    elif ext == ".pdf":
        content = read_pdf(source)
    elif ext in IMAGE_EXTS:
        content = read_image(source)
    else:
        raise SystemExit(f"Extensão ainda não suportada: {ext}")

    out_dir = Path("_extracted")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / f"{source.stem}.txt"
    out_file.write_text(content, encoding="utf-8")

    print(f"Extração concluída: {out_file}")


if __name__ == "__main__":
    main()
